"""
Script to convert COMPASS_APP_DOCUMENTATION.md to a well-formatted .docx file
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def create_docx():
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title Style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 100, 180)
    
    # Heading 1 Style
    h1_style = styles['Heading 1']
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 80, 150)
    
    # Heading 2 Style
    h2_style = styles['Heading 2']
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(50, 50, 50)
    
    # Heading 3 Style
    h3_style = styles['Heading 3']
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(80, 80, 80)
    
    # Read markdown file
    with open('COMPASS_APP_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_content = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_content)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # Skip empty lines at start
        if not line.strip():
            i += 1
            continue
        
        # Title (# )
        if line.startswith('# ') and not line.startswith('## '):
            title_text = line[2:].strip()
            p = doc.add_paragraph(title_text, style='Title')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            i += 1
            continue
        
        # Blockquote (> )
        if line.startswith('> '):
            quote_text = line[2:].strip()
            # Remove markdown formatting
            quote_text = re.sub(r'\*\*([^*]+)\*\*', r'\1', quote_text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(quote_text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue
        
        # Horizontal rule (---)
        if line.strip() == '---':
            doc.add_paragraph('_' * 60)
            i += 1
            continue
        
        # Heading 2 (## )
        if line.startswith('## '):
            heading_text = line[3:].strip()
            doc.add_heading(heading_text, level=1)
            i += 1
            continue
        
        # Heading 3 (### )
        if line.startswith('### '):
            heading_text = line[4:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
            continue
        
        # Heading 4 (#### )
        if line.startswith('#### '):
            heading_text = line[5:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
            continue
        
        # Numbered list
        if re.match(r'^\d+\. ', line):
            list_text = re.sub(r'^\d+\. ', '', line).strip()
            list_text = format_inline_markdown(list_text)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, list_text)
            i += 1
            continue
        
        # Bullet list
        if line.startswith('- '):
            list_text = line[2:].strip()
            list_text = format_inline_markdown(list_text)
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, list_text)
            i += 1
            continue
        
        # Regular paragraph
        para_text = format_inline_markdown(line.strip())
        p = doc.add_paragraph()
        add_formatted_text(p, para_text)
        i += 1
    
    # Save document
    doc.save('COMPASS_APP_DOCUMENTATION.docx')
    print("Successfully created COMPASS_APP_DOCUMENTATION.docx")

def format_inline_markdown(text):
    """Keep markdown markers for processing"""
    return text

def add_formatted_text(paragraph, text):
    """Add text with bold and italic formatting"""
    # Pattern to find **bold** and *italic* text
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|[^*]+)'
    parts = re.findall(pattern, text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Regular text
            paragraph.add_run(part)

if __name__ == '__main__':
    create_docx()
